#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Переделывает билеты в docx:
- все цифры заменяются по перестановке (0→1, 1→2, … 9→0);
- Вася → Катя, У Васи → У Кати, Васей → Катей;
- мальчик → девочка, Мальчик → Девочка.
Составляет матрицу ответов (вариант × задание) для каждого файла.
"""

import re
import csv
from pathlib import Path
from docx import Document
from copy import deepcopy

# Перестановка цифр: каждая цифра заменяется на другую
DIGIT_MAP = str.maketrans('0123456789', '1234567890')

# Защита нумерации вариантов ответа 1)…4) в начале блоков (после \n\n)
# Плейсхолдеры содержат только буквы — не затрагиваются DIGIT_MAP
_OPT_PROTECT = [
    ('\n\n1) ', '\n\n__ОПЦ_ОДИН_ '),
    ('\n\n2) ', '\n\n__ОПЦ_ДВА_ '),
    ('\n\n3) ', '\n\n__ОПЦ_ТРИ_ '),
    ('\n\n4) ', '\n\n__ОПЦ_ЧЕТЫРЕ_ '),
]
_OPT_RESTORE = [(ph, orig) for orig, ph in _OPT_PROTECT]

# Плейсхолдеры для нумерации заданий (без цифр), инструкций и вариантов
TASK_PLACEHOLDERS = [
    ('Задание 10', 'Задание_ДЕСЯТЬ_'),
    ('Задание 9', 'Задание_ДЕВЯТЬ_'),
    ('Задание 8', 'Задание_ВОСЕМЬ_'),
    ('Задание 7', 'Задание_СЕМЬ_'),
    ('Задание 6', 'Задание_ШЕСТЬ_'),
    ('Задание 5', 'Задание_ПЯТЬ_'),
    ('Задание 4', 'Задание_ЧЕТЫРЕ_'),
    ('Задание 3', 'Задание_ТРИ_'),
    ('Задание 2', 'Задание_ДВА_'),
    ('Задание 1', 'Задание_ОДИН_'),
]
_NAMES = ["ОДИН", "ДВА", "ТРИ", "ЧЕТЫРЕ", "ПЯТЬ", "ШЕСТЬ", "СЕМЬ", "ВОСЕМЬ", "ДЕВЯТЬ", "ДЕСЯТЬ"]
TASK_RESTORE = [(f'Задание_{_NAMES[n-1]}_', f'Задание {n}') for n in range(1, 11)]

# Инструкция в начале: номера заданий не трогаем (с пробелом/переносом и без — на случай разбиения по run)
INSTR_PLACEHOLDERS = [
    ('Для заданий 3-7, 9-10 ', 'Для заданий_ИНСТР_ТРИ_СЕМЬ_ДЕВЯТЬ_ДЕСЯТЬ_ '),
    ('Для заданий 3-7, 9-10\n', 'Для заданий_ИНСТР_ТРИ_СЕМЬ_ДЕВЯТЬ_ДЕСЯТЬ_\n'),
    ('Для заданий 3-7, 9-10', 'Для заданий_ИНСТР_ТРИ_СЕМЬ_ДЕВЯТЬ_ДЕСЯТЬ_'),
    ('Для заданий 1, 2, 8 ', 'Для заданий_ИНСТР_ОДИН_ДВА_ВОСЕМЬ_ '),
    ('Для заданий 1, 2, 8\n', 'Для заданий_ИНСТР_ОДИН_ДВА_ВОСЕМЬ_\n'),
    ('Для заданий 1, 2, 8', 'Для заданий_ИНСТР_ОДИН_ДВА_ВОСЕМЬ_'),
    ('Для заданий 2 - 21 ', 'Для заданий_ИНСТР_ДВА_ДВАДЦАТЬ_ОДИН_ '),
    ('Для заданий 2 - 21\n', 'Для заданий_ИНСТР_ДВА_ДВАДЦАТЬ_ОДИН_\n'),
    ('Для заданий 2 - 21', 'Для заданий_ИНСТР_ДВА_ДВАДЦАТЬ_ОДИН_'),
    ('Для заданий 1 - 10 ', 'Для заданий_ИНСТР_ОДИН_ДЕСЯТЬ_ '),
    ('Для заданий 1 - 10\n', 'Для заданий_ИНСТР_ОДИН_ДЕСЯТЬ_\n'),
    ('Для заданий 1 - 10', 'Для заданий_ИНСТР_ОДИН_ДЕСЯТЬ_'),
]
INSTR_RESTORE = [
    ('Для заданий_ИНСТР_ТРИ_СЕМЬ_ДЕВЯТЬ_ДЕСЯТЬ_', 'Для заданий 3-7, 9-10'),
    ('Для заданий_ИНСТР_ОДИН_ДВА_ВОСЕМЬ_', 'Для заданий 1, 2, 8'),
    ('Для заданий_ИНСТР_ДВА_ДВАДЦАТЬ_ОДИН_', 'Для заданий 2 - 21'),
    ('Для заданий_ИНСТР_ОДИН_ДЕСЯТЬ_', 'Для заданий 1 - 10'),
]

# Варианты — плейсхолдеры без цифр (для "Вариант 1" … "Вариант 20" в одном run)
_VAR_DECADE = ["НУЛЬ", "ОДИН", "ДВА", "ТРИ", "ЧЕТЫРЕ", "ПЯТЬ", "ШЕСТЬ", "СЕМЬ", "ВОСЕМЬ", "ДЕВЯТЬ"]
def _var_ph_2(i):
    return f'Вариант_НОМЕР_{_VAR_DECADE[i // 10]}_{_VAR_DECADE[i % 10]}_'
_VAR_NAMES_1_20 = ["ОДИН", "ДВА", "ТРИ", "ЧЕТЫРЕ", "ПЯТЬ", "ШЕСТЬ", "СЕМЬ", "ВОСЕМЬ", "ДЕВЯТЬ", "ДЕСЯТЬ",
                   "ОДИННАДЦАТЬ", "ДВЕНАДЦАТЬ", "ТРИНАДЦАТЬ", "ЧЕТЫРНАДЦАТЬ", "ПЯТНАДЦАТЬ", "ШЕСТНАДЦАТЬ", "СЕМНАДЦАТЬ", "ВОСЕМНАДЦАТЬ", "ДЕВЯТНАДЦАТЬ", "ДВАДЦАТЬ"]
# Обратное отображение: как выглядят номера 01..20 после DIGIT_MAP (12, 13, … 31)
def _variant_mapped(i):
    return f'{i:02d}'.translate(DIGIT_MAP)
_VARIANT_REVERSE = {_variant_mapped(i): f'{i:02d}' for i in range(1, 21)}  # '12'->'01', ...
VARIANT_PLACEHOLDERS = [(f'Вариант № {i:02d}', _var_ph_2(i)) for i in range(20, 0, -1)]
for i in range(20, 0, -1):
    VARIANT_PLACEHOLDERS.append((f'Вариант {i}', f'Вариант_НОМ_{_VAR_NAMES_1_20[i-1]}_'))
VARIANT_RESTORE = [(_var_ph_2(i), f'Вариант № {i:02d}') for i in range(1, 21)]
VARIANT_RESTORE += [(f'Вариант_НОМ_{_VAR_NAMES_1_20[i-1]}_', f'Вариант {i}') for i in range(1, 21)]

def replace_text(s: str) -> str:
    if not s:
        return s
    # 0) Сохраняем инструкции, номера вариантов, нумерацию заданий и год
    for orig, ph in INSTR_PLACEHOLDERS:
        s = s.replace(orig, ph)
    for orig, ph in VARIANT_PLACEHOLDERS:
        s = s.replace(orig, ph)
    for orig, ph in TASK_PLACEHOLDERS:
        s = s.replace(orig, ph)
    s = s.replace('2026', 'ГОД_ДВА_НУЛЬ_ДВА_ШЕСТЬ_')
    s = s.replace('2025', 'ГОД_ДВА_НУЛЬ_ДВА_ШЕСТЬ_')
    s = s.replace('2024', 'ГОД_ДВА_НУЛЬ_ДВА_ШЕСТЬ_')
    # 1) Замена имён (порядок важен: длинные фразы первыми)
    s = s.replace('У Васи ', 'У Кати ')
    s = s.replace('У Васи\n', 'У Кати\n')
    s = s.replace('У Васи.', 'У Кати.')
    s = s.replace('от Васи ', 'от Кати ')
    s = s.replace('от Васи.', 'от Кати.')
    s = s.replace('от Васи по', 'от Кати по')
    s = s.replace('от Васи', 'от Кати')  # любой конец (пробел, запятая и т.д. в др. run)
    s = s.replace('Компьютер Васи ', 'Компьютер Кати ')
    s = s.replace('Васи может', 'Кати может')
    s = s.replace('с Васей ', 'с Катей ')
    s = s.replace('с Васей,', 'с Катей,')
    s = s.replace('Васей ', 'Катей ')
    s = s.replace('Васей\n', 'Катей\n')
    s = s.replace('Васей.', 'Катей.')
    s = s.replace('Вася ', 'Катя ')
    s = s.replace('Вася\n', 'Катя\n')
    s = s.replace('Вася.', 'Катя.')
    s = s.replace('Васи', 'Кати')  # родительный/дательный: компьютер Кати, от Кати
    s = s.replace('Вася', 'Катя')  # оставшиеся
    s = s.replace('мальчик ', 'девочка ')
    s = s.replace('мальчик.', 'девочка.')
    s = s.replace('мальчик\n', 'девочка\n')
    s = s.replace('Мальчик ', 'Девочка ')
    s = s.replace('Мальчик.', 'Девочка.')
    s = s.replace('мальчик', 'девочка')
    s = s.replace('Мальчик', 'Девочка')
    # 2) Замена цифр
    s = s.translate(DIGIT_MAP)
    # 3) Восстанавливаем инструкции, номера вариантов, нумерацию заданий и год 2026
    for ph, orig in INSTR_RESTORE:
        s = s.replace(ph, orig)
    for ph, orig in VARIANT_RESTORE:
        s = s.replace(ph, orig)
    for ph, orig in TASK_RESTORE:
        s = s.replace(ph, orig)
    s = s.replace('ГОД_ДВА_НУЛЬ_ДВА_ШЕСТЬ_', '2026')
    # Год мог быть разбит по run (2025 -> 3136, 2026 -> 3137 после замены цифр) — везде подставляем 2026
    s = s.replace('3136 г.', '2026 г.').replace('3137 г.', '2026 г.')
    s = re.sub(r'\b3136\s*г\.', '2026 г.', s)
    s = re.sub(r'\b3137\s*г\.', '2026 г.', s)
    # Год разбит по run — подставляем 2026 (3136/3137 — типичный результат от 2025)
    s = s.replace('3137', '2026').replace('3136', '2026')
    s = s.replace('5358 г.', '2026 г.').replace('5358', '2026')  # год в «Информационные технологии»
    return s

def process_paragraph(paragraph):
    for run in paragraph.runs:
        if run.text:
            run.text = replace_text(run.text)
    # Номер варианта часто разбит по run («Вариант № », «0», «1») — после замены цифр получается 12, 13… Восстанавливаем 01…20.
    full = paragraph.text
    m = re.search(r'Вариант\s*№\s*(\d+)', full)
    if not m:
        return
    num = m.group(1)
    if num not in _VARIANT_REVERSE:
        return
    correct = _VARIANT_REVERSE[num]
    num_start = full.find(num)
    num_end = num_start + len(num)
    pos = 0
    for run in paragraph.runs:
        rend = pos + len(run.text)
        if num_start < rend and pos < num_end:
            a = max(0, num_start - pos)
            b = min(len(run.text), num_end - pos)
            c_start = a + pos - num_start
            c_end = c_start + (b - a)
            run.text = run.text[:a] + correct[c_start:c_end] + run.text[b:]
        pos = rend

def process_document(doc: Document) -> None:
    for p in doc.paragraphs:
        process_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)
    for section in doc.sections:
        for p in section.header.paragraphs:
            process_paragraph(p)
        for p in section.footer.paragraphs:
            process_paragraph(p)
    # Обработка сносок/концевых сносок, если есть
    try:
        for p in doc.footnotes.paragraphs if hasattr(doc, 'footnotes') and doc.footnotes else []:
            process_paragraph(p)
    except Exception:
        pass

def count_tasks_in_doc(doc: Document) -> int:
    """По первому варианту считаем количество заданий (Задание 1, 2, ...)."""
    tasks = 0
    for p in doc.paragraphs:
        m = re.search(r'Задание\s+(\d+)', p.text, re.IGNORECASE)
        if m:
            tasks = max(tasks, int(m.group(1)))
    return tasks

def get_variant_count(doc: Document) -> int:
    """Количество вариантов по заголовкам «Вариант № XX» или «Вариант X»."""
    count = 0
    for p in doc.paragraphs:
        if re.match(r'Вариант\s*(№\s*)?\d+', p.text.strip(), re.IGNORECASE):
            count += 1
    if count == 0:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if re.search(r'Вариант\s*(№\s*)?\d+', cell.text, re.IGNORECASE):
                        count += 1
    return count

def build_matrix_rows(doc: Document, num_variants: int, num_tasks: int, filename: str,
                     answers_per_variant: list = None) -> list:
    """Строит строки матрицы: одна строка на вариант. answers_per_variant[v] = список ответов для варианта v+1."""
    headers = ['Файл', 'Вариант'] + [f'Задание {i}' for i in range(1, num_tasks + 1)]
    rows = [headers]
    for v in range(1, num_variants + 1):
        variant_num = str(v).zfill(2)  # в матрице и в документе варианты 01, 02, … 20
        ans_row = [''] * num_tasks
        if answers_per_variant and v - 1 < len(answers_per_variant):
            a = answers_per_variant[v - 1]
            for i in range(min(len(a), num_tasks)):
                ans_row[i] = a[i] if a[i] is not None else ''
        row = [filename, variant_num] + ans_row
        rows.append(row)
    return rows


def compute_answers_for_file(doc: Document, num_tasks: int) -> list:
    """По оригинальному документу считает ответы для каждого варианта. Возвращает list of list (вариант -> ответы)."""
    from parse_variant import get_variant_blocks, extract_all_tasks
    from solvers import solve_all
    blocks = get_variant_blocks(doc)
    variant_answers = []
    digit_map_str = '1234567890'
    for start, end in blocks:
        data = extract_all_tasks(doc, start, end)
        ans = solve_all(data, digit_map_str)
        variant_answers.append(ans)
    return variant_answers


def _map_number(x, digit_map_str: str):
    """Применяет перестановку цифр к числу (для ответа «как в переделанном документе»)."""
    if x is None:
        return None
    s = str(x)
    if not digit_map_str:
        return x
    return int(s.translate(str.maketrans('0123456789', digit_map_str))) if s.isdigit() else x


def _map_extracted_data(data: dict, digit_map_str: str) -> dict:
    """Применяет перестановку цифр ко всем числам в извлечённых данных (для решения «как в переделанном документе»)."""
    out = {}
    for k, v in data.items():
        if v is None:
            out[k] = None
        elif k == 1 and isinstance(v, (list, tuple)) and len(v) >= 3:
            opts, low, high = v[0], v[1], v[2]
            opts = list(opts) if opts else []
            mapped_opts = [_map_number(x, digit_map_str) for x in opts]
            out[k] = (mapped_opts, _map_number(low, digit_map_str), _map_number(high, digit_map_str))
        elif k == 2 and isinstance(v, (list, tuple)):
            out[k] = v  # списки строк не трогаем
        elif k == 3 and isinstance(v, (list, tuple)):
            out[k] = v  # кодовые слова — строки из 0/1, можно отобразить посимвольно
        elif k == 4 and isinstance(v, (list, tuple)) and len(v) >= 3:
            out[k] = tuple(_map_number(x, digit_map_str) for x in v[:3])
        elif k == 5 and isinstance(v, (list, tuple)) and len(v) >= 2:
            out[k] = tuple(_map_number(x, digit_map_str) for x in v[:2])
        elif k == 6 and isinstance(v, (list, tuple)) and len(v) >= 3:
            out[k] = tuple(_map_number(x, digit_map_str) if i < 3 else x for i, x in enumerate(v))
        elif k == 7:
            out[k] = v  # задание 7 — без простого отображения
        elif k == 8 and isinstance(v, (list, tuple)) and len(v) == 4:
            p, q, r, opts = v[0], v[1], v[2], v[3]
            mp = (_map_number(p[0], digit_map_str), _map_number(p[1], digit_map_str)) if p else p
            mq = (_map_number(q[0], digit_map_str), _map_number(q[1], digit_map_str)) if q else q
            mr = (_map_number(r[0], digit_map_str), _map_number(r[1], digit_map_str)) if r else r
            out[k] = (mp, mq, mr, [(_map_number(a, digit_map_str), _map_number(b, digit_map_str)) for a, b in opts])
        elif k == 9 and isinstance(v, (list, tuple)) and len(v) >= 2:
            out[k] = tuple(_map_number(v[i], digit_map_str) for i in range(min(4, len(v))))
        elif k == 10 and isinstance(v, int):
            out[k] = _map_number(v, digit_map_str)
        else:
            out[k] = v
    return out


def compute_answers_for_file_mapped(doc: Document, num_tasks: int) -> list:
    """
    Исходные задания не меняем. Парсим оригинал, отображаем параметры перестановкой цифр,
    решаем — ответы соответствуют тому, что видит ученик в переделанном файле.
    """
    from parse_variant import get_variant_blocks, extract_all_tasks
    from solvers import solve_all
    digit_map_str = '1234567890'  # 0→1, 1→2, ... 9→0
    identity_map = '0123456789'
    blocks = get_variant_blocks(doc)
    variant_answers = []
    for start, end in blocks:
        data = extract_all_tasks(doc, start, end)
        mapped_data = _map_extracted_data(data, digit_map_str)
        ans = solve_all(mapped_data, digit_map_trans=identity_map)
        # Задание 3: решали по исходным кодам, ответ нужно отобразить (ученик пишет в отображённых цифрах)
        if len(ans) > 2 and ans[2]:
            ans = list(ans)
            ans[2] = str(ans[2]).translate(DIGIT_MAP)
        variant_answers.append(ans)
    return variant_answers

def _replace_digits_and_names_std(s: str) -> str:
    """Замена цифр и имён для чистых docx (без шапки). Нумерацию 1)…4) в начале блоков сохраняем."""
    if not s:
        return s
    # Защита нумерации вариантов ответа
    for orig, ph in _OPT_PROTECT:
        s = s.replace(orig, ph)
    s = s.replace('Вася', 'Катя').replace('Васи', 'Кати').replace('Васей', 'Катей')
    s = s.replace('мальчик', 'девочка').replace('Мальчик', 'Девочка')
    s = s.translate(DIGIT_MAP)
    # Восстановление нумерации вариантов ответа
    for ph, orig in _OPT_RESTORE:
        s = s.replace(ph, orig)
    return s


def build_clean_standard_doc(base: Path, fname: str, file_id: int) -> tuple:
    """
    Собирает чистый docx без шапки: только Вариант №N и задания №1…№10.
    file_id: 0 = Информатика, 1 = Информационные технологии.
    Возвращает (doc, variant_answers, num_variants=20, num_tasks=10).
    """
    from task_generators import generate_all_for_variant
    doc = Document()
    num_variants, num_tasks = 20, 10
    all_answers = []
    for variant_num in range(1, num_variants + 1):
        effective_variant = variant_num + file_id * 20
        p = doc.add_paragraph()
        p.add_run(f"Вариант №{variant_num}").bold = True
        texts, answers = generate_all_for_variant(effective_variant)
        # Задания с выбором варианта (1, 2, 8) — ответ 1–4, сдвиг цифр не применяем
        _MC = {0, 1, 7}
        answers_mapped = [
            str(a) if i in _MC else str(a).translate(DIGIT_MAP)
            for i, a in enumerate(answers)
        ]
        all_answers.append(answers_mapped)
        for i, text in enumerate(texts):
            text_replaced = _replace_digits_and_names_std(text)
            blocks = [b.strip() for b in text_replaced.split("\n\n") if b.strip()]
            doc.add_paragraph(f"№{i + 1}")
            for block in blocks[1:]:  # первый блок — «Задание N.» пропускаем
                doc.add_paragraph(block)
    return doc, all_answers, num_variants, num_tasks


def main():
    base = Path(__file__).parent
    # Два файла со стандартными заданиями (20 вариантов × 10 заданий), без шапки
    standard_files = [
        ('Информатика.docx', 0),
        ('Информационные технологии.docx', 1),
    ]

    all_matrix_rows = []
    num_tasks_per_file = {}
    file_answers = {}

    for fname, file_id in standard_files:
        print(f'Обработка: {fname}')
        doc, variant_answers, num_variants, num_tasks = build_clean_standard_doc(base, fname, file_id)
        num_tasks_per_file[fname] = (num_variants, num_tasks)
        file_answers[fname] = variant_answers
        out_name = Path(fname).stem + '_переделанные.docx'
        out_path = base / out_name
        doc.save(out_path)
        print(f'  Сохранено: {out_name}, вариантов: {num_variants}, заданий: {num_tasks}')

        rows = build_matrix_rows(doc, num_variants, num_tasks, fname, answers_per_variant=variant_answers)
        if not all_matrix_rows:
            all_matrix_rows = [rows[0]]
        all_matrix_rows.extend(rows[1:])

    # Матрица в один CSV (все файлы)
    matrix_path = base / 'матрица_ответов.csv'
    with open(matrix_path, 'w', newline='', encoding='utf-8-sig') as f:
        w = csv.writer(f, delimiter=';')
        w.writerows(all_matrix_rows)
    print(f'Матрица ответов: {matrix_path.name}')
    # Пояснение, почему не все ячейки заполнены
    expl_path = base / 'Почему не все ответы посчитаны.txt'
    with open(expl_path, 'w', encoding='utf-8') as f:
        f.write("""Матрица ответов

Два файла без шапки: только Вариант №N и задания №1…№10 (Информатика, Информационные технологии). Файл для собеседования (Информатика 2) — build_interview_docs.py.
""")
    print(f'  Пояснение: {expl_path.name}')

    # Отдельная матрица по каждому файлу
    for fname, _ in standard_files:
        if fname not in num_tasks_per_file:
            continue
        num_variants, num_tasks = num_tasks_per_file[fname]
        doc = Document()  # для матрицы нужны только размеры и ответы
        rows = build_matrix_rows(doc, num_variants, num_tasks, fname,
                                 answers_per_variant=file_answers.get(fname))
        one_path = base / (Path(fname).stem + '_матрица_ответов.csv')
        with open(one_path, 'w', newline='', encoding='utf-8-sig') as f:
            w = csv.writer(f, delimiter=';')
            w.writerows(rows)
        print(f'  Матрица по файлу: {one_path.name}')

    # Ответы для собеседования: один файл с ответами по всем вариантам для быстрой проверки
    write_answer_key_for_interview(base, file_answers, num_tasks_per_file, [f for f, _ in standard_files])


def write_answer_key_for_interview(base: Path, file_answers: dict, num_tasks_per_file: dict, files: list):
    """Пишет «Ответы_для_собеседования.txt» и «Ответы_для_собеседования.docx» — удобно при проведении собеседования."""
    lines = []
    lines.append("ОТВЕТЫ ДЛЯ СОБЕСЕДОВАНИЯ")
    lines.append("20 вариантов × 10 заданий. Задания в каждом варианте разные (разные числа и сценарии).")
    lines.append("")
    for fname in files:
        if fname not in num_tasks_per_file or fname not in file_answers:
            continue
        num_variants, num_tasks = num_tasks_per_file[fname]
        answers = file_answers[fname]
        doc_name = Path(fname).stem + '_переделанные.docx'
        lines.append("=" * 60)
        lines.append(f"Файл: {doc_name}")
        lines.append("=" * 60)
        headers = "Вариант | " + " | ".join(f"Зад.{i}" for i in range(1, num_tasks + 1))
        lines.append(headers)
        lines.append("-" * 60)
        for v in range(num_variants):
            a = answers[v] if v < len(answers) else []
            row = [str(v + 1).zfill(2)]
            for i in range(num_tasks):
                row.append(a[i] if i < len(a) and a[i] is not None else "—")
            lines.append("   " + "   |   ".join(row))
        lines.append("")
    txt_path = base / "Ответы_для_собеседования.txt"
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(lines))
    print(f'Ответы для собеседования: {txt_path.name}')

    # Docx с таблицей ответов по каждому файлу
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        doc.add_paragraph("Ответы для собеседования", style='Title')
        doc.add_paragraph("20 вариантов × 10 заданий. Используйте при проверке ответов на собеседовании.")
        for fname in files:
            if fname not in num_tasks_per_file or fname not in file_answers:
                continue
            num_variants, num_tasks = num_tasks_per_file[fname]
            answers = file_answers[fname]
            doc_name = Path(fname).stem + '_переделанные.docx'
            doc.add_paragraph(doc_name, style='Heading 2')
            table = doc.add_table(rows=1 + num_variants, cols=1 + num_tasks)
            table.style = 'Table Grid'
            h = table.rows[0].cells
            h[0].text = "Вариант"
            for i in range(num_tasks):
                h[i + 1].text = f"Зад.{i + 1}"
            for v in range(num_variants):
                row = table.rows[v + 1].cells
                row[0].text = str(v + 1)
                a = answers[v] if v < len(answers) else []
                for i in range(num_tasks):
                    row[i + 1].text = a[i] if i < len(a) and a[i] is not None else "—"
            doc.add_paragraph("")
        docx_path = base / "Ответы_для_собеседования.docx"
        doc.save(docx_path)
        print(f'  Таблица ответов: {docx_path.name}')
    except Exception as e:
        print(f'  Docx ответов не создан: {e}')


if __name__ == '__main__':
    main()
