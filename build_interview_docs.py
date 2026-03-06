#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Один файл с заданиями для устного собеседования (без шапок):
- Информатика 2: 20 вариантов, в каждом 5 вопросов с развёрнутым ответом + 5 заданий (разные числа).
- Файл ответов: эталоны на вопросы + ответы на задания.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm

from interview_questions import THEORY_QUESTIONS, THEORY_ANSWERS
from task_generators import gen_task1, gen_task2, gen_task3, gen_task4, gen_task5

# Перестановка цифр 0→1..9→0 (как в основных билетах)
DIGIT_MAP = str.maketrans('0123456789', '1234567890')

# Плейсхолдеры для нумерации заданий и вопросов (чтобы не портить заменой цифр)
_TASK_PH = [(f'Задание {n}.', f'Задание_НОМ_{n}_') for n in range(5, 0, -1)]  # 5,4,3,2,1 — длинные первыми
_TASK_RESTORE = [(f'Задание_НОМ_{n}_', f'Задание {n}.') for n in range(1, 6)]

# Защита нумерации вариантов ответа 1)…4) в начале блоков (после \n\n)
_OPT_PROTECT = [
    ('\n\n1) ', '\n\n__ОПЦ_ОДИН_ '),
    ('\n\n2) ', '\n\n__ОПЦ_ДВА_ '),
    ('\n\n3) ', '\n\n__ОПЦ_ТРИ_ '),
    ('\n\n4) ', '\n\n__ОПЦ_ЧЕТЫРЕ_ '),
]
_OPT_RESTORE = [(ph, orig) for orig, ph in _OPT_PROTECT]


def replace_digits_and_names(s: str) -> str:
    """Замена цифр и имён. Нумерацию «Задание 1.» … «Задание 5.» и варианты 1)…4) сохраняем."""
    if not s:
        return s
    for orig, ph in _TASK_PH:
        s = s.replace(orig, ph)
    # Защита нумерации вариантов ответа
    for orig, ph in _OPT_PROTECT:
        s = s.replace(orig, ph)
    s = s.replace('Вася', 'Катя').replace('Васи', 'Кати').replace('Васей', 'Катей')
    s = s.replace('мальчик', 'девочка').replace('Мальчик', 'Девочка')
    s = s.translate(DIGIT_MAP)
    for ph, orig in _TASK_RESTORE:
        s = s.replace(ph, orig)
    # Восстановление нумерации вариантов ответа
    for ph, orig in _OPT_RESTORE:
        s = s.replace(ph, orig)
    return s


# Один файл для собеседования (file_id=2 → другие числа, чем в стандартных Информатика / Информационные технологии)
DISCIPLINES = [
    ('Информатика 2', 2),
]

TASK_GENS = [gen_task1, gen_task2, gen_task3, gen_task4, gen_task5]
NUM_VARIANTS = 20
NUM_QUESTIONS = 5
NUM_TASKS = 5


def _add_variant_header(doc, text: str):
    """Добавляет заголовок варианта: жирный, 14pt, с отступом снизу."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    return p


def _add_section_label(doc, text: str):
    """Добавляет подзаголовок раздела (курсив, жирный)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(11)
    return p


def _add_task_block(doc, number: str, blocks: list):
    """Добавляет номер задания и его содержимое с отступом."""
    p_num = doc.add_paragraph()
    p_num.paragraph_format.space_before = Pt(6)
    p_num.paragraph_format.space_after = Pt(2)
    r = p_num.add_run(number)
    r.bold = True
    r.font.size = Pt(11)
    for block in blocks:
        p = doc.add_paragraph(block)
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)


def build_one_docx(discipline_name: str, file_id: int, base: Path) -> list:
    """
    Создаёт один docx: только варианты, без шапок.
    Возвращает список ответов по вариантам: [ [задание1, ..., задание5], ... ]
    """
    doc = Document()

    # Поля страницы
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    all_answers = []

    for variant_num in range(1, NUM_VARIANTS + 1):
        effective_variant = variant_num + file_id * 20

        # Разрыв страницы перед каждым вариантом, кроме первого
        if variant_num > 1:
            doc.add_page_break()

        _add_variant_header(doc, f"Вариант №{variant_num}")

        # №1–№5: теоретические вопросы
        _add_section_label(doc, "Теоретические вопросы:")
        for i in range(NUM_QUESTIONS):
            _add_task_block(doc, f"№{i + 1}", [THEORY_QUESTIONS[i]])

        # №6–№10: задания
        _add_section_label(doc, "Задания:")
        task_answers = []
        for i, gen in enumerate(TASK_GENS):
            text, ans = gen(effective_variant)
            text_replaced = replace_digits_and_names(text)
            blocks = [b.strip() for b in text_replaced.split("\n\n") if b.strip()]
            _add_task_block(doc, f"№{NUM_QUESTIONS + i + 1}", blocks[1:])
            # Задания с выбором варианта (1, 2) — ответ 1–4, сдвиг цифр не применяем
            _MC = {0, 1}
            if i in _MC:
                task_answers.append(str(ans))
            else:
                task_answers.append(str(ans).translate(DIGIT_MAP))
        all_answers.append(task_answers)

    out_name = f"{discipline_name}_собеседование.docx"
    doc.save(base / out_name)
    return all_answers


def write_answer_file(base: Path, discipline_answers: dict):
    """Пишет файл с ответами: теоретические вопросы один раз + компактная таблица ответов."""
    lines = []
    lines.append("ОТВЕТЫ ДЛЯ УСТНОГО СОБЕСЕДОВАНИЯ")
    lines.append(f"Структура: {NUM_QUESTIONS} теоретических вопросов + {NUM_TASKS} заданий")
    lines.append("")

    for discipline_name, file_id in DISCIPLINES:
        key = (discipline_name, file_id)
        if key not in discipline_answers:
            continue
        task_answers_per_variant = discipline_answers[key]

        lines.append("=" * 70)
        lines.append(f"  Дисциплина: {discipline_name}")
        lines.append(f"  Файл билетов: {discipline_name}_собеседование.docx")
        lines.append("=" * 70)
        lines.append("")

        # Теоретические вопросы — ОДИН РАЗ для всех вариантов
        lines.append("─" * 70)
        lines.append("  ЭТАЛОНЫ ОТВЕТОВ НА ТЕОРЕТИЧЕСКИЕ ВОПРОСЫ")
        lines.append("  (вопросы одинаковы во всех вариантах)")
        lines.append("─" * 70)
        for i in range(NUM_QUESTIONS):
            lines.append("")
            lines.append(f"  Вопрос {i + 1}.  {THEORY_QUESTIONS[i]}")
            lines.append(f"  Ответ:    {THEORY_ANSWERS[i]}")
        lines.append("")

        # Компактная таблица ответов на задания
        lines.append("─" * 70)
        lines.append("  ОТВЕТЫ НА ЗАДАНИЯ (краткие)")
        lines.append("─" * 70)
        col = 10
        header = f"  {'Вариант':^8}" + "".join(f"{'Зад.' + str(i + 1):^{col}}" for i in range(NUM_TASKS))
        lines.append(header)
        lines.append("  " + "─" * (8 + col * NUM_TASKS))
        for v in range(NUM_VARIANTS):
            ans = task_answers_per_variant[v]
            row = f"  {v + 1:^8}" + "".join(
                f"{(ans[i] if i < len(ans) else '—'):^{col}}" for i in range(NUM_TASKS)
            )
            lines.append(row)
        lines.append("")

    txt_path = base / "Ответы_собеседование.txt"
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    # Docx с таблицей ответов на задания (для быстрой проверки)
    doc = Document()
    doc.add_paragraph("Ответы для устного собеседования", style="Title")
    doc.add_paragraph("Эталоны на вопросы — в файле Ответы_собеседование.txt. В таблице ниже — только ответы на задания (1–5).")
    doc.add_paragraph("")
    for discipline_name, file_id in DISCIPLINES:
        key = (discipline_name, file_id)
        if key not in discipline_answers:
            continue
        task_answers_per_variant = discipline_answers[key]
        doc.add_paragraph(f"{discipline_name}_собеседование.docx", style="Heading 2")
        table = doc.add_table(rows=1 + NUM_VARIANTS, cols=1 + NUM_TASKS)
        table.style = "Table Grid"
        h = table.rows[0].cells
        h[0].text = "Вариант"
        for i in range(NUM_TASKS):
            h[i + 1].text = f"Зад.{i + 1}"
        for v in range(NUM_VARIANTS):
            row = table.rows[v + 1].cells
            row[0].text = str(v + 1)
            for i in range(NUM_TASKS):
                row[i + 1].text = task_answers_per_variant[v][i]
        doc.add_paragraph("")
    docx_path = base / "Ответы_собеседование.docx"
    doc.save(docx_path)

    return txt_path, docx_path


def main():
    base = Path(__file__).resolve().parent
    discipline_answers = {}

    for discipline_name, file_id in DISCIPLINES:
        print(f"Формирование: {discipline_name}_собеседование.docx")
        answers = build_one_docx(discipline_name, file_id, base)
        discipline_answers[(discipline_name, file_id)] = answers
        print(f"  Вариантов: {NUM_VARIANTS}, вопросов: {NUM_QUESTIONS}, заданий: {NUM_TASKS}")

    txt_path, docx_path = write_answer_file(base, discipline_answers)
    print(f"Ответы (эталоны + задания): {txt_path.name}")
    print(f"Таблица ответов на задания: {docx_path.name}")


if __name__ == "__main__":
    main()
