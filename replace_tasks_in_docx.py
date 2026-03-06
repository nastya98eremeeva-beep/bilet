# -*- coding: utf-8 -*-
"""
Замена содержимого заданий в docx на сгенерированные (со 100% считающимися ответами).
Возвращает variant_answers: список из 20 списков по 10 ответов.
Поддержка документов, где контент в таблицах (в т.ч. Информатика 2).
"""

import re
from docx import Document
from docx.oxml.ns import qn
from parse_variant import get_variant_blocks, block_text
from task_generators import generate_all_for_variant


def _all_paragraphs_ordered(doc: Document) -> list:
    """Все параграфы в порядке появления в документе (включая параграфы внутри таблиц)."""
    result = []
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            for p in doc.paragraphs:
                if p._element is child:
                    result.append(p)
                    break
        elif child.tag == qn('w:tbl'):
            for table in doc.tables:
                if table._tbl is child:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                result.append(p)
                    break
    return result


def get_task_boundaries(doc, start_para: int, end_para: int) -> list:
    """Для одного варианта возвращает [(start, end), ...] для заданий 1..10 (только doc.paragraphs)."""
    task_starts = []
    for i in range(start_para, end_para):
        if i >= len(doc.paragraphs):
            break
        m = re.match(r'Задание\s+(\d+)', doc.paragraphs[i].text.strip(), re.I)
        if m:
            task_starts.append((i, int(m.group(1))))
    if len(task_starts) < 10:
        return []
    boundaries = []
    for j in range(10):
        a = task_starts[j][0]
        b = task_starts[j + 1][0] if j + 1 < len(task_starts) else end_para
        boundaries.append((a, b))
    return boundaries


def _task_heading_match(text: str) -> int:
    """Возвращает номер задания (1-10) если параграф — заголовок задания, иначе 0."""
    t = (text or '').strip()
    m = re.match(r'Задание\s+(\d+)', t, re.I)
    if m:
        return int(m.group(1))
    m = re.match(r'^\s*(\d+)\.\s', t)
    if m:
        n = int(m.group(1))
        if 1 <= n <= 10:
            return n
    return 0


def _replace_tasks_with_flat_paras(doc: Document, all_paras: list, file_id: int) -> list:
    """Ищет варианты и задания в плоском списке параграфов (в т.ч. из таблиц). Возвращает variant_answers."""
    variant_starts = []
    for i, p in enumerate(all_paras):
        if re.match(r'Вариант\s*(№\s*)?\d+', (p.text or '').strip(), re.I):
            variant_starts.append(i)
    if not variant_starts:
        return []
    variant_answers = []
    for v_idx in range(len(variant_starts)):
        start = variant_starts[v_idx]
        end = variant_starts[v_idx + 1] if v_idx + 1 < len(variant_starts) else len(all_paras)
        seen_num = {}
        for i in range(start, end):
            num = _task_heading_match(all_paras[i].text)
            if num and 1 <= num <= 10 and num not in seen_num:
                seen_num[num] = i
        indices = [seen_num.get(j) for j in range(1, 11)]
        effective_variant = (v_idx + 1) + file_id * 20
        texts, answers = generate_all_for_variant(effective_variant)
        variant_answers.append(answers)
        if all(indices):
            for j in range(10):
                a, b = indices[j], indices[j + 1] if j + 1 < 10 else end
                all_paras[a].text = texts[j]
                for i in range(a + 1, b):
                    if i < len(all_paras):
                        all_paras[i].text = ''
        else:
            # Нет 10 заголовков заданий — подменяем весь блок варианта: вариант + 10 параграфов под задачи
            need = start + 1 + 10
            for j in range(10):
                idx = start + 1 + j
                if idx < len(all_paras):
                    all_paras[idx].text = texts[j]
            for i in range(start + 11, end):
                if i < len(all_paras):
                    all_paras[i].text = ''
    return variant_answers


def _file_id_from_name(filename: str) -> int:
    """Числовой идентификатор файла для разной генерации заданий (вариативность между файлами)."""
    if not filename:
        return 0
    name = filename.lower().replace('.docx', '').strip()
    if 'информационные технологии' in name or 'технологии' in name:
        return 1
    if 'информатика 2' in name or 'информатика2' in name:
        return 2
    return 0  # Информатика.docx и прочие


def replace_tasks_in_document(doc: Document, filename: str = None) -> list:
    """
    Заменяет текст заданий 1–10 во всех вариантах на сгенерированный.
    filename — имя файла (для вариативности: в разных файлах разные задания при одном номере варианта).
    Возвращает variant_answers: list из num_variants списков по 10 ответов (строки).
    Если границы заданий не найдены по doc.paragraphs (например контент в таблицах), используется обход всех параграфов включая таблицы.
    """
    file_id = _file_id_from_name(filename or '')
    blocks = get_variant_blocks(doc)
    variant_answers = []
    use_flat = not blocks
    for v_idx, (start, end) in enumerate(blocks):
        variant_num = v_idx + 1
        effective_variant = variant_num + file_id * 20
        boundaries = get_task_boundaries(doc, start, end)
        if len(boundaries) != 10:
            use_flat = True
            break
        texts, answers = generate_all_for_variant(effective_variant)
        variant_answers.append(answers)
        for task_idx in range(10):
            a, b = boundaries[task_idx]
            doc.paragraphs[a].text = texts[task_idx]
            for i in range(a + 1, b):
                if i < len(doc.paragraphs):
                    doc.paragraphs[i].text = ''
    if use_flat and len(variant_answers) < len(blocks) if blocks else True:
        all_paras = _all_paragraphs_ordered(doc)
        if all_paras:
            variant_answers = _replace_tasks_with_flat_paras(doc, all_paras, file_id)
    return variant_answers
